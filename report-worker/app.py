import os
import mysql.connector
from docx import Document
import requests

DB_SERVER = os.getenv("OCS_DB_SERVER", "ocsinventory-db")
DB_USER = os.getenv("OCS_DB_USER", "ocsuser")
DB_PASS = os.getenv("OCS_DB_PASS", "ocspass")
DB_NAME = os.getenv("OCS_DB_NAME", "ocsweb")

OLLAMA_API_URL = "http://ollama:11434/api/chat"

def query_ocs_data():
    conn = mysql.connector.connect(
        host=DB_SERVER,
        user=DB_USER,
        password=DB_PASS,
        database=DB_NAME
    )
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT name, ip_addr, last_seen FROM hardware ORDER BY last_seen DESC LIMIT 20")
    rows = cursor.fetchall()
    cursor.close()
    conn.close()
    return rows

def generate_report_text(ocs_data):
    devices_list = "\n".join([f"- {d['name']} ({d['ip_addr']}) visto: {d['last_seen']}" for d in ocs_data])
    prompt = f"Genera un report formale intitolato 'Assessment infrastrutturale IT' basato su questi dati di inventario:\n{devices_list}\n\nFornisci una sintesi e raccomandazioni."

    payload = {
        "model": "mistral",
        "messages": [{"role": "user", "content": prompt}]
    }

    response = requests.post(OLLAMA_API_URL, json=payload)
    response.raise_for_status()
    result = response.json()
    return result['choices'][0]['message']['content']

def create_docx_report(text, filename="Assessment_Infrastrutturale_IT.docx"):
    doc = Document()
    doc.add_heading("Assessment infrastrutturale IT", 0)
    doc.add_paragraph(text)
    doc.save(filename)

if __name__ == "__main__":
    data = query_ocs_data()
    ai_text = generate_report_text(data)
    create_docx_report(ai_text)
    print(f"Report generato in Assessment_Infrastrutturale_IT.docx")
